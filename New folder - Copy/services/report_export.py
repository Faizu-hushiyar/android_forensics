"""Export case report text to PDF and Word (.docx)."""

from __future__ import annotations

import io
import re


def _plain_from_markdown(md: str) -> str:
    t = re.sub(r"^#+\s*", "", md, flags=re.MULTILINE)
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"[*_`]", "", t)
    return t


def export_pdf_bytes(text: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    for line in text.split("\n"):
        if not line.strip():
            pdf.ln(5)
            continue
            
        if line.startswith("### "):
            pdf.set_font("Helvetica", "B", 14)
            text_to_print = line.replace("### ", "").replace("**", "")
            pdf.ln(2)
        elif line.startswith("#### "):
            pdf.set_font("Helvetica", "B", 12)
            text_to_print = line.replace("#### ", "").replace("**", "")
            pdf.ln(2)
        elif line.startswith("---"):
            pdf.set_font("Helvetica", "", 10)
            text_to_print = "_" * 60
        else:
            pdf.set_font("Helvetica", "", 10)
            # Remove markdown bold/code ticks, but preserve underscores for file paths
            text_to_print = line.replace("**", "").replace("`", "")
            
        # Ensure characters are printable in standard latin-1 Helvetica
        safe = "".join(c if ord(c) < 256 else "?" for c in text_to_print)
        
        try:
            # write handles wrapping and avoids fpdf2 multi_cell crash on unspaced strings
            pdf.write(5, safe + "\n")
        except Exception:
            pdf.write(5, "[Error rendering line]\n")
            
    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return out.encode("latin-1", errors="replace")


def export_docx_bytes(text: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Security Analytics — Case Report", 0)
    
    for block in text.split("\n\n"):
        if not block.strip():
            continue
            
        lines = block.splitlines()
        for line in lines:
            line = line.strip()
            if line.startswith("### "):
                doc.add_heading(line.replace("### ", "").replace("**", ""), level=2)
            elif line.startswith("#### "):
                doc.add_heading(line.replace("#### ", "").replace("**", ""), level=3)
            elif line.startswith("---"):
                doc.add_paragraph("_" * 60)
            else:
                plain = line.replace("**", "").replace("`", "").replace("_", "")
                doc.add_paragraph(plain[:8000] or " ")
        doc.add_paragraph("")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
